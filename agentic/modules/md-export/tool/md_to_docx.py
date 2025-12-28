#!/usr/bin/env python3
"""
Convert Markdown files to formatted Word documents (.docx).

Uses python-docx for full control over document styling, including:
- Callout boxes with dotted borders
- Clean heading hierarchy
- Native tables
- Mermaid diagrams rendered as embedded images
- Preserved flowcharts as monospace blocks

Usage:
    # Single file
    ./run execution/md_to_docx.py path/to/file.md

    # Custom output path
    ./run execution/md_to_docx.py path/to/file.md --output custom.docx

    # Directory (recursive)
    ./run execution/md_to_docx.py path/to/directory/

    # Preview without creating
    ./run execution/md_to_docx.py path/to/file.md --dry-run

Requirements:
    - python-docx (pip install python-docx)
    - mermaid-cli for diagram rendering (npm install -g @mermaid-js/mermaid-cli)
"""

import sys
import os
import re
import json
import argparse
import hashlib
import subprocess
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ==============================================================================
# Mermaid Diagram Rendering
# ==============================================================================

# Cache directory for rendered diagrams
MERMAID_CACHE_DIR = Path(__file__).parent.parent / '.tmp' / 'mermaid'


def get_mermaid_cli_path() -> Optional[str]:
    """Find the mermaid-cli (mmdc) executable."""
    # Check if mmdc is in PATH
    mmdc_path = shutil.which('mmdc')
    if mmdc_path:
        return mmdc_path

    # Common installation locations
    common_paths = [
        '/usr/local/bin/mmdc',
        '/opt/homebrew/bin/mmdc',
        os.path.expanduser('~/.npm-global/bin/mmdc'),
        os.path.expanduser('~/node_modules/.bin/mmdc'),
    ]

    for path in common_paths:
        if os.path.isfile(path) and os.access(path, os.X_OK):
            return path

    return None


def render_mermaid(content: str) -> Optional[Path]:
    """Render Mermaid diagram to PNG.

    Args:
        content: Mermaid diagram source code

    Returns:
        Path to rendered PNG, or None if rendering failed
    """
    mmdc_path = get_mermaid_cli_path()
    if not mmdc_path:
        print("  Warning: mermaid-cli (mmdc) not found. Install with: npm install -g @mermaid-js/mermaid-cli", file=sys.stderr)
        return None

    # Create cache directory
    MERMAID_CACHE_DIR.mkdir(parents=True, exist_ok=True)

    # Generate hash for caching
    content_hash = hashlib.md5(content.encode()).hexdigest()[:12]

    # Check cache
    png_path = MERMAID_CACHE_DIR / f"{content_hash}.png"
    if png_path.exists():
        print(f"  Using cached diagram: {content_hash}", file=sys.stderr)
        return png_path

    # Write Mermaid source to temp file
    mmd_path = MERMAID_CACHE_DIR / f"{content_hash}.mmd"
    with open(mmd_path, 'w', encoding='utf-8') as f:
        f.write(content)

    # Render with mmdc
    try:
        result = subprocess.run(
            [
                mmdc_path,
                '-i', str(mmd_path),
                '-o', str(png_path),
                '-b', 'white',  # White background for Word docs
                '-s', '2',      # Scale factor for higher resolution
            ],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode != 0:
            print(f"  Warning: Mermaid rendering failed: {result.stderr}", file=sys.stderr)
            return None

        if png_path.exists():
            print(f"  Rendered diagram: {content_hash}", file=sys.stderr)
            return png_path
        else:
            print("  Warning: Mermaid output file not created", file=sys.stderr)
            return None

    except subprocess.TimeoutExpired:
        print("  Warning: Mermaid rendering timed out", file=sys.stderr)
        return None
    except Exception as e:
        print(f"  Warning: Mermaid rendering error: {e}", file=sys.stderr)
        return None


# ==============================================================================
# ASCII Art Detection
# ==============================================================================

BOX_CHARS = r'[┌┐└┘│═─╔╗╚╝║├┤┬┴┼┃┏┓┗┛]'
FLOW_CHARS = r'[▼▲→←↓↑]'
BRANCH_PATTERN = r'[/\\]\s+[/\\]'


def is_flowchart(content: str) -> bool:
    """Check if content is a flowchart (should be preserved as code block)."""
    return bool(re.search(FLOW_CHARS, content) or re.search(BRANCH_PATTERN, content))


def is_ascii_box(content: str) -> bool:
    """Check if content contains ASCII box characters."""
    return bool(re.search(BOX_CHARS, content))


def extract_box_content(content: str) -> Tuple[Optional[str], List[str]]:
    """Extract title and items from an ASCII box.

    Returns:
        (title, items) tuple where title may be None
    """
    lines = content.split('\n')
    title = None
    items = []

    for line in lines:
        # Remove box-drawing characters
        clean = re.sub(BOX_CHARS, '', line).strip()

        if not clean:
            continue

        # First non-empty line is title
        if title is None:
            title = clean
            continue

        # Process remaining lines
        if clean.startswith('✓') or clean.startswith('✔'):
            items.append(('check', clean.lstrip('✓✔ ').strip()))
        elif clean.startswith('•'):
            items.append(('bullet', clean.lstrip('• ').strip()))
        elif clean.startswith('❌') or clean.startswith('✗'):
            items.append(('x', clean.lstrip('❌✗ ').strip()))
        elif clean.endswith(':') and len(clean) < 50:
            items.append(('subheader', clean))
        elif clean.isupper() and len(clean) > 10 and ':' not in clean:
            items.append(('tagline', clean))
        else:
            items.append(('bullet', clean))

    return title, items


# ==============================================================================
# Markdown Parsing
# ==============================================================================

def parse_markdown(content: str) -> List[Dict[str, Any]]:
    """Parse markdown into a list of block elements.

    Returns list of dicts with 'type' and type-specific fields.
    """
    blocks = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith('```'):
            # Check for language specifier (e.g., ```mermaid, ```python)
            lang_match = re.match(r'^```(\w+)?', stripped)
            lang = lang_match.group(1) if lang_match else None

            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_content = '\n'.join(code_lines)

            # Handle Mermaid diagrams specially
            if lang and lang.lower() == 'mermaid':
                blocks.append({'type': 'mermaid', 'content': code_content})
            elif is_flowchart(code_content):
                blocks.append({'type': 'code_block', 'content': code_content})
            elif is_ascii_box(code_content):
                title, items = extract_box_content(code_content)
                blocks.append({'type': 'callout_box', 'title': title, 'items': items})
            else:
                blocks.append({'type': 'code_block', 'content': code_content})
            i += 1
            continue

        # Horizontal rule
        if stripped in ['---', '***', '___']:
            blocks.append({'type': 'horizontal_rule'})
            i += 1
            continue

        # Heading
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                blocks.append({'type': 'heading', 'level': level, 'text': text})
                i += 1
                continue

        # Table
        if '|' in stripped and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            # Parse table
            rows = []
            for tline in table_lines:
                cells = [c.strip() for c in tline.strip('|').split('|')]
                # Skip separator row
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue
                rows.append(cells)

            if rows:
                blocks.append({'type': 'table', 'rows': rows})
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('> '))
                i += 1
            blocks.append({'type': 'blockquote', 'text': '\n'.join(quote_lines)})
            continue

        # Bullet list
        if re.match(r'^[-*+]\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^[-*+]\s', lines[i].strip()):
                item_text = re.sub(r'^[-*+]\s+', '', lines[i].strip())
                items.append(item_text)
                i += 1
            blocks.append({'type': 'bullet_list', 'items': items})
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                items.append(item_text)
                i += 1
            blocks.append({'type': 'numbered_list', 'items': items})
            continue

        # Paragraph (non-empty lines)
        if stripped:
            para_lines = []
            while i < len(lines):
                current = lines[i].strip()
                if not current:
                    break
                # Check for block-level elements that should stop paragraph
                if current.startswith('#'):
                    break
                if current.startswith('```'):
                    break
                if current.startswith('>'):
                    break
                if '|' in current and i + 1 < len(lines) and '|' in lines[i + 1]:
                    break
                # Check for list items (must have space after marker)
                if re.match(r'^[-*+]\s', current):
                    break
                if re.match(r'^\d+\.\s', current):
                    break
                # Check for horizontal rule
                if current in ['---', '***', '___']:
                    break
                para_lines.append(current)
                i += 1
            if para_lines:
                blocks.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
            continue

        # Skip empty lines
        i += 1

    return blocks


# ==============================================================================
# Document Building
# ==============================================================================

def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex: str):
    """Set background shading for a paragraph."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), color_hex)
    paragraph._p.get_or_add_pPr().append(shading)


def add_border_to_paragraph(paragraph, border_color: str = '999999', border_style: str = 'dotted'):
    """Add a border around a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    border_val = 'dotted' if border_style == 'dotted' else 'single'

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), border_val)
        border.set(qn('w:sz'), '6')  # Border size
        border.set(qn('w:space'), '4')  # Space between border and text
        border.set(qn('w:color'), border_color)
        pBdr.append(border)

    pPr.append(pBdr)


def apply_inline_formatting(paragraph, text: str):
    """Apply inline formatting (bold, italic, code) to text."""
    # Pattern for **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        full_match = match.group(0)
        if full_match.startswith('**'):
            # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif full_match.startswith('`'):
            # Code
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def convert_checkmarks(text: str) -> str:
    """Convert checkbox syntax and Unicode checkmarks to emojis."""
    text = re.sub(r'\[x\]', '✅', text, flags=re.IGNORECASE)
    text = re.sub(r'\[ \]', '⬜', text)
    text = text.replace('✓', '✅')
    text = text.replace('✔', '✅')
    text = text.replace('☑', '✅')
    text = text.replace('☒', '✅')
    text = text.replace('☐', '⬜')
    text = text.replace('✗', '❌')
    return text


def create_document() -> Document:
    """Create a new document with custom styles."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    return doc


def add_heading(doc: Document, text: str, level: int):
    """Add a heading with appropriate styling."""
    text = convert_checkmarks(text)

    heading = doc.add_heading(level=level)
    apply_inline_formatting(heading, text)

    # Style adjustments based on level
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_paragraph_text(doc: Document, text: str):
    """Add a regular paragraph with inline formatting."""
    text = convert_checkmarks(text)
    para = doc.add_paragraph()
    apply_inline_formatting(para, text)


def add_bullet_list(doc: Document, items: List[str]):
    """Add a bullet list."""
    for item in items:
        item = convert_checkmarks(item)
        para = doc.add_paragraph(style='List Bullet')
        apply_inline_formatting(para, item)


def add_numbered_list(doc: Document, items: List[str]):
    """Add a numbered list."""
    for item in items:
        item = convert_checkmarks(item)
        para = doc.add_paragraph(style='List Number')
        apply_inline_formatting(para, item)


def add_table(doc: Document, rows: List[List[str]]):
    """Add a native Word table."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = convert_checkmarks(cell_text)

                # Header row styling
                if i == 0:
                    set_cell_shading(cell, 'E8E8E8')
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    # Add spacing after table
    doc.add_paragraph()


def add_blockquote(doc: Document, text: str):
    """Add a blockquote with italic styling and indentation."""
    text = convert_checkmarks(text)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    run = para.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_code_block(doc: Document, content: str):
    """Add a code block with monospace font and gray background."""
    para = doc.add_paragraph()
    set_paragraph_shading(para, 'F0F0F0')
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)

    run = para.add_run(content)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)


def add_callout_box(doc: Document, title: Optional[str], items: List[Tuple[str, str]]):
    """Add a callout box with dotted border."""
    # Create the callout paragraph
    if title:
        title_para = doc.add_paragraph()
        set_paragraph_shading(title_para, 'F5F5F5')
        add_border_to_paragraph(title_para, '999999', 'dotted')

        # Add title with emoji if it looks like a rules/tips box
        display_title = title
        title_upper = title.upper()
        if 'RULE' in title_upper or 'TIP' in title_upper:
            display_title = f"📋 {title}"
        elif 'QUICK' in title_upper or 'REF' in title_upper:
            display_title = f"📋 {title}"
        elif 'WARNING' in title_upper:
            display_title = f"⚠️ {title}"
        elif 'NOTE' in title_upper:
            display_title = f"📝 {title}"

        run = title_para.add_run(display_title)
        run.bold = True
        run.font.size = Pt(12)

    # Add items
    for item_type, item_text in items:
        para = doc.add_paragraph()
        set_paragraph_shading(para, 'F5F5F5')
        add_border_to_paragraph(para, '999999', 'dotted')
        para.paragraph_format.left_indent = Inches(0.25)

        if item_type == 'check':
            para.add_run('✅ ')
            apply_inline_formatting(para, item_text)
        elif item_type == 'x':
            para.add_run('❌ ')
            apply_inline_formatting(para, item_text)
        elif item_type == 'subheader':
            run = para.add_run(item_text)
            run.bold = True
        elif item_type == 'tagline':
            run = para.add_run(item_text)
            run.bold = True
            run.font.size = Pt(11)
        else:
            para.add_run('• ')
            apply_inline_formatting(para, item_text)

    # Spacing after callout
    doc.add_paragraph()


def add_horizontal_rule(doc: Document):
    """Add a page break or spacing for section separation."""
    # Add some spacing (could be changed to page break if desired)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)


def get_image_dimensions(image_path: Path) -> Tuple[int, int]:
    """Get image dimensions without PIL dependency.

    Returns (width, height) in pixels by reading PNG header.
    """
    try:
        with open(image_path, 'rb') as f:
            # PNG signature + IHDR chunk
            f.read(8)  # Skip PNG signature
            f.read(4)  # Skip chunk length
            f.read(4)  # Skip 'IHDR'
            width = int.from_bytes(f.read(4), 'big')
            height = int.from_bytes(f.read(4), 'big')
            return width, height
    except Exception:
        return 0, 0


def add_mermaid_diagram(doc: Document, content: str):
    """Render and embed a Mermaid diagram.

    Falls back to code block if rendering fails.
    Automatically sizes image to fit within page margins.
    """
    png_path = render_mermaid(content)

    if png_path and png_path.exists():
        # Add centered paragraph with image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Get image dimensions to calculate appropriate size
        # Default Word page is 8.5" x 11" with 1" margins = 6.5" usable width
        MAX_WIDTH_INCHES = 5.5  # Leave some breathing room
        MAX_HEIGHT_INCHES = 7.0  # Don't let tall diagrams dominate

        img_width, img_height = get_image_dimensions(png_path)

        if img_width > 0 and img_height > 0:
            # Scale factor is 2 (from mmdc -s 2), so actual dimensions are half
            # But we work with the rendered size for aspect ratio
            aspect_ratio = img_width / img_height

            # Start with max width and calculate height
            width_inches = MAX_WIDTH_INCHES
            height_inches = width_inches / aspect_ratio

            # If too tall, constrain by height instead
            if height_inches > MAX_HEIGHT_INCHES:
                height_inches = MAX_HEIGHT_INCHES
                width_inches = height_inches * aspect_ratio

            # For smaller/simpler diagrams, don't stretch them unnecessarily
            # If the natural size (at 96 DPI) is smaller, use that
            natural_width = img_width / 96 / 2  # Divide by scale factor
            if natural_width < width_inches and natural_width >= 2.0:
                width_inches = natural_width

            run = para.add_run()
            run.add_picture(str(png_path), width=Inches(width_inches))
        else:
            # Fallback if we can't read dimensions
            run = para.add_run()
            run.add_picture(str(png_path), width=Inches(4.5))

        # Add spacing after diagram
        doc.add_paragraph()
    else:
        # Fallback: show as code block with note
        para = doc.add_paragraph()
        para.add_run("[Mermaid diagram - install mermaid-cli to render]").italic = True

        add_code_block(doc, content)


# ==============================================================================
# Main Conversion
# ==============================================================================

def convert_markdown_to_docx(md_path: Path, output_path: Path = None, dry_run: bool = False) -> Dict[str, Any]:
    """Convert a markdown file to a Word document.

    Args:
        md_path: Path to input markdown file
        output_path: Optional output path (defaults to same name with .docx)
        dry_run: If True, parse but don't create document

    Returns:
        Dict with result info (path, title, etc.)
    """
    print(f"Processing: {md_path}", file=sys.stderr)

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Parse markdown
    blocks = parse_markdown(content)

    # Extract title from first heading
    title = md_path.stem.replace('-', ' ').replace('_', ' ').title()
    for block in blocks:
        if block['type'] == 'heading' and block['level'] == 1:
            title = block['text']
            break

    print(f"  Title: {title}", file=sys.stderr)
    print(f"  Blocks: {len(blocks)}", file=sys.stderr)

    if dry_run:
        print(f"  [DRY RUN] Would create: {title}.docx", file=sys.stderr)
        return {
            'title': title,
            'file': str(md_path),
            'blocks': len(blocks),
            'dry_run': True
        }

    # Determine output path
    if output_path is None:
        output_path = md_path.with_suffix('.docx')

    # Create document
    doc = create_document()

    # Build document from blocks
    # Stop at appendix section (contains diagram source code for editing, not for client docs)
    APPENDIX_MARKERS = ['appendix: diagram sources', 'appendix: diagram source']

    for block in blocks:
        block_type = block['type']

        # Check if we've reached the appendix section
        if block_type == 'heading':
            heading_lower = block['text'].lower().strip()
            if any(marker in heading_lower for marker in APPENDIX_MARKERS):
                print("  Excluding appendix section from Word output", file=sys.stderr)
                break  # Stop processing - don't include appendix in Word doc

        if block_type == 'heading':
            add_heading(doc, block['text'], block['level'])
        elif block_type == 'paragraph':
            add_paragraph_text(doc, block['text'])
        elif block_type == 'bullet_list':
            add_bullet_list(doc, block['items'])
        elif block_type == 'numbered_list':
            add_numbered_list(doc, block['items'])
        elif block_type == 'table':
            add_table(doc, block['rows'])
        elif block_type == 'blockquote':
            add_blockquote(doc, block['text'])
        elif block_type == 'code_block':
            add_code_block(doc, block['content'])
        elif block_type == 'mermaid':
            add_mermaid_diagram(doc, block['content'])
        elif block_type == 'callout_box':
            add_callout_box(doc, block['title'], block['items'])
        elif block_type == 'horizontal_rule':
            add_horizontal_rule(doc)

    # Save document
    doc.save(str(output_path))

    print(f"  Created: {output_path}", file=sys.stderr)

    return {
        'path': str(output_path),
        'title': title,
        'blocks': len(blocks)
    }


def process_directory(dir_path: Path, dry_run: bool = False) -> List[Dict[str, Any]]:
    """Process all markdown files in a directory."""
    results = []

    md_files = sorted(dir_path.rglob('*.md'))

    if not md_files:
        print(f"No markdown files found in {dir_path}", file=sys.stderr)
        return results

    print(f"Found {len(md_files)} markdown file(s)", file=sys.stderr)

    for md_path in md_files:
        try:
            result = convert_markdown_to_docx(md_path, dry_run=dry_run)
            results.append(result)
        except Exception as e:
            print(f"  Error: {e}", file=sys.stderr)
            results.append({
                'file': str(md_path),
                'error': str(e)
            })

    return results


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to formatted Word documents'
    )
    parser.add_argument(
        'path',
        help='Path to Markdown file or directory'
    )
    parser.add_argument(
        '--output', '-o',
        help='Output file path (for single file conversion)'
    )
    parser.add_argument(
        '--dry-run', '-n',
        action='store_true',
        help='Parse and show what would be created without creating'
    )

    args = parser.parse_args()

    path = Path(args.path)

    if not path.exists():
        print(f"Error: Path not found: {path}", file=sys.stderr)
        sys.exit(1)

    if path.is_file():
        if path.suffix != '.md':
            print(f"Error: Not a markdown file: {path}", file=sys.stderr)
            sys.exit(1)

        output_path = Path(args.output) if args.output else None
        result = convert_markdown_to_docx(path, output_path, args.dry_run)
        print(json.dumps(result, indent=2))
    else:
        results = process_directory(path, args.dry_run)
        print(json.dumps(results, indent=2))


if __name__ == '__main__':
    main()
