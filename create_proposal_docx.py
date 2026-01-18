from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_line(doc, text, level=1):
    """Add heading with blue underline"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return heading

doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============ COVER PAGE ============
# Blue header bar (simulated with paragraph)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(40)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI OPPORTUNITY\nASSESSMENT")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Divider
divider = doc.add_paragraph("—————")
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Proposal
prop = doc.add_paragraph()
prop.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prop.add_run("Proposal")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

doc.add_paragraph()

# Prepared for
prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prep.add_run("Prepared for\n")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
run = prep.add_run("Sandy Bomyea, S&S Wolf Sheds")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# Presented by
pres = doc.add_paragraph()
pres.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pres.add_run("Presented by\n")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
run = pres.add_run("Arise AI Group\n")
run.bold = True
run.font.size = Pt(14)
run = pres.add_run("arisegroup.ai")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

doc.add_page_break()

# ============ PAGE 2: THE OPPORTUNITY ============
h = doc.add_heading("The Opportunity", level=1)

p = doc.add_paragraph()
p.add_run("Every organization has processes that drain time, introduce errors, and limit growth. The challenge isn't knowing AI can help — it's knowing ")
p.add_run("where to start").bold = True
p.add_run(" and ")
p.add_run("which opportunities will deliver real ROI").bold = True
p.add_run(".")

p = doc.add_paragraph("This AI Opportunity Assessment gives you clarity. In 2-3 weeks, you'll move from uncertainty to a validated roadmap — with prioritized opportunities, quantified business impact, and a clear path to implementation.")

# What You Get
doc.add_heading("What You Get", level=1)
doc.add_paragraph("This engagement delivers actionable outcomes, not just recommendations:")

# Deliverables table
deliverables = [
    ("Prioritized AI Opportunities", "A clear view of where AI will have the greatest impact — mapped by business value and implementation effort so you know exactly what to tackle first."),
    ("Quantified ROI Projections", "Data-backed business cases for each opportunity, including projected cost savings and revenue impact — the numbers to make confident decisions."),
    ("Implementation Roadmap", "A sequenced plan showing what to implement first, dependencies, and realistic timeframes — move from assessment to action immediately."),
    ("Executive Presentation", "A polished deliverable ready for leadership alignment — the complete story from current challenges to recommended solutions."),
]

table = doc.add_table(rows=len(deliverables), cols=1)
table.style = 'Table Grid'
for i, (title, desc) in enumerate(deliverables):
    cell = table.rows[i].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(title + "\n")
    run.bold = True
    p.add_run(desc)
    set_cell_shading(cell, "F5F5F5")

doc.add_paragraph()

# Quick Wins section
p = doc.add_paragraph()
run = p.add_run("Plus Quick Wins Delivered During Assessment:")
run.bold = True

quick_wins = [
    ("Initial SOPs", "Documented processes for lead handling and lot operations — capturing your team's knowledge in repeatable, trainable formats."),
    ("Website Quick Fixes", "Broken images repaired, mobile issues fixed, load times improved — professional first impression restored."),
    ("Lead Qualification System", "Interactive \"Find which shed is best for you\" questionnaire replacing basic contact form — pre-qualify leads before they consume your time."),
]

table2 = doc.add_table(rows=len(quick_wins), cols=1)
table2.style = 'Table Grid'
for i, (title, desc) in enumerate(quick_wins):
    cell = table2.rows[i].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(title + "\n")
    run.bold = True
    p.add_run(desc)
    set_cell_shading(cell, "E8F4E8")

doc.add_page_break()

# ============ PAGE 3: OUR APPROACH ============
doc.add_heading("Our Approach", level=1)

p = doc.add_paragraph()
p.add_run("The foundation of a successful AI initiative is understanding your business from the inside out. Our goal is to become an ")
run = p.add_run("\"inefficiency detective\"")
run.bold = True
p.add_run(" — uncovering the hidden friction, repetitive tasks, and manual processes that drain time and resources.")

# Discovery Interviews
doc.add_heading("Discovery Interviews", level=2)

p = doc.add_paragraph()
p.add_run("We conduct targeted interviews with two key groups: ")
p.add_run("leadership").bold = True
p.add_run(" to understand strategic goals and how success is measured, and ")
p.add_run("end-users").bold = True
p.add_run(" to understand the on-the-ground reality of daily tasks. The gap between these two perspectives is often where the biggest opportunities lie.")

# Interview table for SS Wolf Sheds
interview_table = doc.add_table(rows=5, cols=3)
interview_table.style = 'Table Grid'
headers = ["Who", "Duration", "What We Learn"]
for i, h in enumerate(headers):
    cell = interview_table.rows[0].cells[i]
    cell.paragraphs[0].add_run(h).bold = True
    set_cell_shading(cell, "4472C4")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

interviews = [
    ("Sandy (Owner)", "60 min", "Revenue, pricing, lead flow, close rate"),
    ("Matthew (Operations)", "60 min", "Operating costs, CODB validation"),
    ("Alex (Website)", "60 min", "Traffic data, inquiry patterns"),
    ("Scott (Transportation)", "60 min", "Delivery data, build error frequency"),
]
for i, (who, dur, what) in enumerate(interviews):
    interview_table.rows[i+1].cells[0].text = who
    interview_table.rows[i+1].cells[1].text = dur
    interview_table.rows[i+1].cells[2].text = what

doc.add_paragraph()

# Process Mapping
doc.add_heading("Process Mapping (Ops Canvas)", level=2)

doc.add_paragraph("Interview insights are translated into a visual Ops Canvas that maps your core operations. For each process, we identify and tag friction points:")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Time Sinks: ").bold = True
p.add_run("Tasks that are highly manual, repetitive, and consume significant employee time")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Quality Risks: ").bold = True
p.add_run("Steps prone to human error or inconsistencies")

doc.add_paragraph("These tagged items are your goldmine — every single one represents a potential opportunity for an AI solution.")

# Add Ops Canvas image
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sample Ops Canvas")
run.italic = True

try:
    doc.add_picture("c:/Users/MEKAIEL/Downloads/proposal_images/image1.png", width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph("[Ops Canvas Image]")

doc.add_page_break()

# ============ PAGE 4: OPPORTUNITY PRIORITIZATION ============
doc.add_heading("Opportunity Prioritization", level=2)

p = doc.add_paragraph()
p.add_run("Each identified pain point is evaluated and plotted on our Opportunity Matrix based on two dimensions: ")
p.add_run("Business Impact").bold = True
p.add_run(" and ")
p.add_run("Implementation Effort").bold = True
p.add_run(". This creates an instant visualization of which AI projects to tackle first.")

# Add Opportunity Matrix image
try:
    doc.add_picture("c:/Users/MEKAIEL/Downloads/proposal_images/image2.png", width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph("[Opportunity Matrix Image]")

doc.add_paragraph()

# Quadrant explanations
quad_table = doc.add_table(rows=2, cols=2)
quad_table.style = 'Table Grid'

quads = [
    ("Quick Wins", "Low effort, high impact. Your #1 priority — fast value, builds trust and momentum."),
    ("Big Swings", "High effort, high impact. Transformative projects for after quick wins are in place."),
    ("Nice-to-Haves", "Low effort, low impact. Minor efficiency gains — good for extra value, not the focus."),
    ("Deprioritize", "High effort, low impact. Time/money pits to avoid — identifying these is part of our value."),
]

for i, (title, desc) in enumerate(quads):
    row = i // 2
    col = i % 2
    cell = quad_table.rows[row].cells[col]
    p = cell.paragraphs[0]
    run = p.add_run(title + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
    p.add_run(desc)

doc.add_paragraph()

# Collaborative Validation
doc.add_heading("Collaborative Validation", level=2)
doc.add_paragraph("The Opportunity Matrix is a powerful hypothesis, but it's not final until validated against your deep internal knowledge. We schedule a follow-up session to review the matrix together as a collaborative workshop — by the end, we'll be presenting a plan that we've already agreed on together.")

doc.add_page_break()

# ============ PAGE 5: ENGAGEMENT OVERVIEW ============
doc.add_heading("Engagement Overview", level=1)
doc.add_paragraph("This focused engagement runs 2-3 weeks and delivers both the strategic assessment and immediate quick wins you can see working right away.")

# Phases table
phases_table = doc.add_table(rows=4, cols=2)
phases_table.style = 'Table Grid'

phases = [
    ("Phase 1\nScoping", "Collaborative session to identify and align on the highest-priority processes for assessment."),
    ("Phase 2\nDiscovery", "Stakeholder and end-user interviews (~60 min each) to understand workflows, pain points, and opportunities. Process mapping to visualize operations and identify time sinks and quality risks."),
    ("Phase 3\nQuick Wins", "Website fixes deployed, lead qualification system live, initial SOPs documented — tangible value delivered during the assessment."),
    ("Phase 4\nSolution Design", "Opportunity prioritization, ROI analysis, roadmap development, and collaborative validation with your team before final executive presentation."),
]

for i, (phase, desc) in enumerate(phases):
    cell0 = phases_table.rows[i].cells[0]
    cell0.paragraphs[0].add_run(phase).bold = True
    set_cell_shading(cell0, "4472C4")
    cell0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    phases_table.rows[i].cells[1].text = desc

doc.add_paragraph()

# Investment
doc.add_heading("Investment", level=1)

inv_table = doc.add_table(rows=1, cols=1)
inv_table.style = 'Table Grid'
cell = inv_table.rows[0].cells[0]
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("AI Opportunity Assessment\n").bold = True
run = p.add_run("$5,000\n")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
p.add_run("2-3 week engagement  •  Includes quick wins delivery")
set_cell_shading(cell, "F5F5F5")

doc.add_page_break()

# ============ PAGE 6: WHY ARISE AI GROUP ============
doc.add_heading("Why Arise AI Group", level=1)
doc.add_paragraph("Arise AI Group is an AI Transformation Partner helping organizations convert manual work into revenue-generating systems. We don't just build AI — we help you identify where AI will have the greatest impact and guide you through implementation.")

p = doc.add_paragraph()
run = p.add_run("What makes us different:")
run.bold = True

bullets = [
    ("Interview first", "Understand YOUR business before building"),
    ("Validate together", "Co-create priorities, not dictate them"),
    ("You own everything", "Database accounts, documents, all of it"),
    ("Quick wins included", "Tangible value delivered during assessment, not just a report"),
]

for title, desc in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + " — ").bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

# CTA Box
cta_table = doc.add_table(rows=1, cols=1)
cta_table.style = 'Table Grid'
cell = cta_table.rows[0].cells[0]
set_cell_shading(cell, "2F3E46")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ready to Get Started?\n")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(255, 255, 255)
run = p.add_run("Let's schedule a scoping session to identify your priority areas.\n")
run.font.color.rgb = RGBColor(255, 255, 255)
run = p.add_run("arisegroup.ai")
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)

# Save
output_path = "c:/Users/MEKAIEL/Downloads/SS_Wolf_Sheds_AI_Opportunity_Assessment.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
